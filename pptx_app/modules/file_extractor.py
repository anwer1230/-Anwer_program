"""
استخراج النصوص والجداول والصور من ملفات Word وPDF
"""

import os
import io
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple


def extract_from_word(file_bytes: bytes) -> Dict[str, Any]:
    """استخراج كل المحتوى من ملف Word"""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    result = {
        "text_blocks": [],
        "tables": [],
        "images": [],
        "full_text": "",
    }

    image_dir = Path(tempfile.mkdtemp())
    image_index = 0

    # استخراج الصور من العلاقات
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                ext = rel.target_part.content_type.split("/")[-1]
                if ext in ("jpeg", "jpg", "png", "gif", "bmp", "tiff"):
                    img_path = image_dir / f"img_{image_index}.{ext}"
                    img_path.write_bytes(img_bytes)
                    result["images"].append(str(img_path))
                    image_index += 1
            except Exception:
                pass

    # استخراج الفقرات
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        result["text_blocks"].append({"type": "paragraph", "style": style, "text": text})
        result["full_text"] += text + "\n"

    # استخراج الجداول
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            result["tables"].append(rows)

    return result


def extract_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
    """استخراج كل المحتوى من ملف PDF"""
    result = {
        "text_blocks": [],
        "tables": [],
        "images": [],
        "full_text": "",
    }

    try:
        import pdfplumber

        image_dir = Path(tempfile.mkdtemp())
        image_index = 0

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # نص
                text = page.extract_text() or ""
                if text.strip():
                    result["text_blocks"].append(
                        {"type": "page", "style": f"صفحة {page_num}", "text": text.strip()}
                    )
                    result["full_text"] += text.strip() + "\n"

                # جداول
                for tbl in page.extract_tables():
                    if tbl:
                        cleaned = [
                            [cell or "" for cell in row] for row in tbl
                        ]
                        result["tables"].append(cleaned)

                # صور - عبر fitz إن توفرت
                try:
                    import fitz  # PyMuPDF

                    doc_fitz = fitz.open(stream=file_bytes, filetype="pdf")
                    page_fitz = doc_fitz[page_num - 1]
                    for img_info in page_fitz.get_images():
                        xref = img_info[0]
                        base_image = doc_fitz.extract_image(xref)
                        img_bytes = base_image["image"]
                        ext = base_image["ext"]
                        img_path = image_dir / f"pdf_img_{image_index}.{ext}"
                        img_path.write_bytes(img_bytes)
                        result["images"].append(str(img_path))
                        image_index += 1
                    doc_fitz.close()
                except Exception:
                    pass

    except Exception as e:
        result["full_text"] = f"خطأ في قراءة PDF: {str(e)}"

    return result


def extract_content(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """دالة موحدة لاستخراج المحتوى"""
    ext = Path(filename).suffix.lower()
    if ext in (".docx", ".doc"):
        return extract_from_word(file_bytes)
    elif ext == ".pdf":
        return extract_from_pdf(file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
        return {
            "text_blocks": [{"type": "paragraph", "style": "Normal", "text": text}],
            "tables": [],
            "images": [],
            "full_text": text,
        }
